import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from .analyse_data import descriptive, plot_line_series, plot_box_plot

def generar_informe(data, tickers, weights):
    doc = Document()
    doc.add_heading('Informe de Análisis de Inversión', level=1)
    doc.add_paragraph("Este informe contiene un análisis de los pesos óptimos del portafolio, "
                      "estadísticas descriptivas de los activos y gráficos visuales de su rendimiento.")

    doc.add_heading('Estadísticas Descriptivas', level=2)
    stats = descriptive(data)
    table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'

    for i, column in enumerate(stats.columns):
        hdr_cells[i + 1].text = column

    for index, row in stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = index
        for i, value in enumerate(row):
            row_cells[i + 1].text = str(value)

    doc.add_heading('Gráficos de Rendimiento', level=2)

    fig1_path = "temp_plot_line_series.png"
    plot_line_series(data, tickers)
    plt.savefig(fig1_path, format='png', bbox_inches='tight')
    plt.close()  
    doc.add_paragraph("Gráfico de Series Temporales:")
    doc.add_picture(fig1_path, width=Inches(6))

    fig2_path = "temp_plot_box_plot.png"
    plot_box_plot(data, tickers)
    plt.savefig(fig2_path, format='png', bbox_inches='tight')
    plt.close()  
    doc.add_paragraph("Gráfico de Box Plot:")
    doc.add_picture(fig2_path, width=Inches(6))

    doc.add_heading('Pesos Óptimos del Portafolio', level=2)
    for ticker, weight in weights.items():
        doc.add_paragraph(f"{ticker}: {float(weight):.2%}")

    if os.path.exists(fig1_path):
        os.remove(fig1_path)
    if os.path.exists(fig2_path):
        os.remove(fig2_path)

    doc_path = 'informe_analisis_inversion.docx'
    doc.save(doc_path)
    file_path = os.path.abspath(doc_path)
    print(f"El informe de análisis ha sido guardado como '{doc_path}' en '{file_path}'")